
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
import os


def create_word_doc_title(doc_title = ''):
    CURR_DIR = os.getcwd()
    print(CURR_DIR)
    doc = Document(r"./is_abg/custom_apps/doc/Intersight.AsBuilt.docx")
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    doc.add_heading(doc_title, 0)
    p = doc.add_paragraph('[Insert Introdution Here]')
    p.keep_together=True
    
    return doc


def create_word_doc_paragraph(doc, heading_text = '', heading_level = 1,
                            paragraph_text = ''):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    doc.add_page_break()
    doc.add_heading(heading_text, level=heading_level)
    p = doc.add_paragraph(paragraph_text)
    p.keep_together=True

    return doc


def create_word_doc_table(doc, df):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'List Table 4 Accent 1')

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    doc.add_page_break()
    return doc